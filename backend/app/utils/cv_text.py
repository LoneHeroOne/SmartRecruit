"""
Text extraction utilities for CV/resume processing
"""
import os
from typing import str
from pypdf import PdfReader
from docx import Document


def extract_text_from_file(file_path: str) -> str:
    """
    Extract text from various file formats (PDF, DOCX, TXT)

    Args:
        file_path: Path to the file

    Returns:
        Extracted text content

    Raises:
        ValueError: If file type is unsupported
        FileNotFoundError: If file doesn't exist
        Exception: For other extraction errors
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    file_extension = file_path.lower()

    try:
        if file_extension.endswith('.pdf'):
            return _extract_pdf_text(file_path)
        elif file_extension.endswith('.docx'):
            return _extract_docx_text(file_path)
        elif file_extension.endswith('.txt'):
            return _extract_txt_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {os.path.basename(file_path)}. "
                           "Supported types: PDF, DOCX, TXT")
    except Exception as e:
        raise Exception(f"Failed to extract text from {file_path}: {str(e)}")


def _extract_pdf_text(file_path: str) -> str:
    """Extract text from PDF file"""
    reader = PdfReader(file_path)
    text_parts = []

    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)

    return "\n".join(text_parts)


def _extract_docx_text(file_path: str) -> str:
    """Extract text from DOCX file"""
    doc = Document(file_path)
    text_parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Also extract from tables if present
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)

    return "\n".join(text_parts)


def _extract_txt_text(file_path: str) -> str:
    """Extract text from TXT file"""
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()


def clean_extracted_text(text: str) -> str:
    """
    Clean and normalize extracted text

    Args:
        text: Raw extracted text

    Returns:
        Cleaned text
    """
    if not text:
        return ""

    # Remove excessive whitespace and normalize line breaks
    text = " ".join(text.split())

    # Remove non-printable characters but keep basic punctuation
    import re
    text = re.sub(r'[^\w\s.,!?-]', ' ', text)

    # Normalize spaces
    text = " ".join(text.split())

    return text.strip()


def get_file_info(file_path: str) -> dict:
    """
    Get basic information about a file

    Args:
        file_path: Path to the file

    Returns:
        Dictionary with file information
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    stat = os.stat(file_path)
    file_extension = os.path.splitext(file_path)[1].lower()

    return {
        'path': file_path,
        'filename': os.path.basename(file_path),
        'extension': file_extension,
        'size_bytes': stat.st_size,
        'size_kb': round(stat.st_size / 1024, 2),
        'modified_time': stat.st_mtime
    }


def preview_text(text: str, max_length: int = 500) -> str:
    """
    Get a preview of text, truncated to max_length

    Args:
        text: Full text
        max_length: Maximum length for preview

    Returns:
        Text preview
    """
    if not text:
        return ""

    if len(text) <= max_length:
        return text

    return text[:max_length].strip() + "..."
